from vaderSentiment.vaderSentiment import SentimentIntensityAnalyzer
import pandas as pd
import docx
import numpy as np
import csv
import pandas as pd

def test():
    sentences=[]

    document = docx.Document("testdata/Dataset.docx")

    for para in document.paragraphs:
        sentences.append(para.text)

    paragraphSentiments = 0.0
    analyzer = SentimentIntensityAnalyzer()
    for sentence in sentences:
        vs = analyzer.polarity_scores(sentence)
        print("{:-<65} {}".format(sentence, str(vs)))
    print("AVERAGE SENTIMENT FOR PARAGRAPH: \t" + str(round(paragraphSentiments / len(sentences), 4)))

def test_news():
    sentences = []

    document = docx.Document("testdata/Dataset_news.docx")

    for para in document.paragraphs:
        sentences.append(para.text)

    paragraphSentiments = 0.0
    analyzer = SentimentIntensityAnalyzer()

    positive=[]
    negative=[]
    neutral=[]

    for sentence in sentences:
        vs = analyzer.polarity_scores(sentence)
        print("{:-<65} {}".format(sentence, str(vs)))
        positive.append(vs["pos"])
        negative.append(vs["neg"])
        neutral.append(vs["neu"])
    print("AVERAGE SENTIMENT FOR PARAGRAPH: \t" + str(round(paragraphSentiments / len(sentences), 4)))
    df=pd.DataFrame({"neg":negative,
                     "neu":neutral,
                     "pos":positive})
    df.to_csv("testdata/news_tri.csv")

def test_story():
    sentences = []

    document = docx.Document("testdata/Dataset_story.docx")

    positive = []
    negative = []
    neutral = []

    for para in document.paragraphs:
        sentences.append(para.text)

    paragraphSentiments = 0.0
    analyzer = SentimentIntensityAnalyzer()
    for sentence in sentences:
        vs = analyzer.polarity_scores(sentence)
        positive.append(vs["pos"])
        negative.append(vs["neg"])
        neutral.append(vs["neu"])
        print("{:-<65} {}".format(sentence, str(vs)))
    print("AVERAGE SENTIMENT FOR PARAGRAPH: \t" + str(round(paragraphSentiments / len(sentences), 4)))
    df = pd.DataFrame({"neg": negative,
                       "neu": neutral,
                       "pos": positive})
    df.to_csv("testdata/story_tri.csv")

def test_scifi():
    sentences = []

    document = docx.Document("testdata/Dataset_sci-fi.docx")

    positive = []
    negative = []
    neutral = []

    for para in document.paragraphs:
        sentences.append(para.text)

    paragraphSentiments = 0.0
    analyzer = SentimentIntensityAnalyzer()
    for sentence in sentences:
        vs = analyzer.polarity_scores(sentence)
        positive.append(vs["pos"])
        negative.append(vs["neg"])
        neutral.append(vs["neu"])
        print("{:-<65} {}".format(sentence, str(vs)))
    print("AVERAGE SENTIMENT FOR PARAGRAPH: \t" + str(round(paragraphSentiments / len(sentences), 4)))
    df = pd.DataFrame({"neg": negative,
                       "neu": neutral,
                       "pos": positive})
    df.to_csv("testdata/sci-fi_tri.csv")

if __name__ == "__main__":
    print('-----------------------------------------------------------------Test Dataset-----------------------------------------------------------------')
    test()
    print(
        '-----------------------------------------------------------------News Dataset-----------------------------------------------------------------')
    test_news()
    print(
        '-----------------------------------------------------------------Story Dataset-----------------------------------------------------------------')
    test_story()
    print(
        '-----------------------------------------------------------------Sci-fi Dataset-----------------------------------------------------------------')
    test_scifi()
